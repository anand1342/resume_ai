<<<<<<< HEAD
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from docx import Document
from pypdf import PdfReader
import os
import uuid

from resume_parser import parse_resume
from gap_analyzer import build_timeline, detect_gaps
from resume_optimizer import generate_resume
from exporter import export_to_word

app = FastAPI()
templates = Jinja2Templates(directory="templates")

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/generate")
async def generate_resume_web(
    request: Request,
    target_role: str = Form(...),
    file: UploadFile = File(...)
):
    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_FOLDER, file_id + "_" + file.filename)

    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    ext = os.path.splitext(file.filename)[1].lower()

    if ext == ".txt":
        original_resume = open(file_path, encoding="utf-8", errors="ignore").read()
    elif ext in [".doc", ".docx"]:
        doc = Document(file_path)
        original_resume = "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".pdf":
        reader = PdfReader(file_path)
        original_resume = "\n".join(p.extract_text() or "" for p in reader.pages)
    else:
        return {"error": "Unsupported file type"}

    parsed = parse_resume(original_resume)
    timeline = build_timeline(parsed)
    gaps = detect_gaps(timeline)

    final_resume = generate_resume(
        target_role,
        original_resume,
        parsed["education"],
        parsed["employment"],
        []
    )

    output_file = f"Generated_{file_id}.docx"
    export_to_word(final_resume, output_file)

    return FileResponse(
        output_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="ATS_Resume.docx"
    )
=======
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from docx import Document
from pypdf import PdfReader
import os
import uuid

from resume_parser import parse_resume
from gap_analyzer import build_timeline, detect_gaps
from resume_optimizer import generate_resume
from exporter import export_to_word

app = FastAPI()
templates = Jinja2Templates(directory="templates")

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/generate")
async def generate_resume_web(
    request: Request,
    target_role: str = Form(...),
    file: UploadFile = File(...)
):
    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_FOLDER, file_id + "_" + file.filename)

    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    ext = os.path.splitext(file.filename)[1].lower()

    if ext == ".txt":
        original_resume = open(file_path, encoding="utf-8", errors="ignore").read()
    elif ext in [".doc", ".docx"]:
        doc = Document(file_path)
        original_resume = "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".pdf":
        reader = PdfReader(file_path)
        original_resume = "\n".join(p.extract_text() or "" for p in reader.pages)
    else:
        return {"error": "Unsupported file type"}

    parsed = parse_resume(original_resume)
    timeline = build_timeline(parsed)
    gaps = detect_gaps(timeline)

    final_resume = generate_resume(
        target_role,
        original_resume,
        parsed["education"],
        parsed["employment"],
        []
    )

    output_file = f"Generated_{file_id}.docx"
    export_to_word(final_resume, output_file)

    return FileResponse(
        output_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="ATS_Resume.docx"
    )
>>>>>>> 35c7c4d302eceb9c4ef79c15ff7b8d631289440e
