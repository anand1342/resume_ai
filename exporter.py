<<<<<<< HEAD
from docx import Document

def export_to_word(text, filename):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
=======
from docx import Document

def export_to_word(text, filename):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
>>>>>>> 35c7c4d302eceb9c4ef79c15ff7b8d631289440e
